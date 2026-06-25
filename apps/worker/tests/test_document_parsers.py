from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document as DocxDocument

from parsers import docling_parser
from parsers import parse_file, parse_file_to_document
from parsers.artifact import ParsedDocument, ParserInfo, ParseQuality, build_blocks_from_output


DOCX_MAIN_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
DOTX_MAIN_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"


def test_parse_txt_decodes_utf8_with_replacement(tmp_path):
    path = tmp_path / "defense.txt"
    path.write_bytes("Gate 2 text ".encode("utf-8") + b"\xff" + " metrics".encode("utf-8"))

    parsed = parse_file(path)

    assert parsed == "Gate 2 text \ufffd metrics"


def test_parse_markdown_returns_source_text(tmp_path):
    path = tmp_path / "defense.md"
    path.write_text("# Gate 2\n\nMVP scope and risks.", encoding="utf-8")

    parsed = parse_file(path)

    assert parsed == "# Gate 2\n\nMVP scope and risks."


def test_parse_markdown_returns_structured_artifact_blocks(tmp_path):
    path = tmp_path / "defense.md"
    path.write_text("# Gate 2\n\nMVP scope and risks.", encoding="utf-8")

    parsed = parse_file_to_document(path)

    assert parsed.plain_text == "# Gate 2\n\nMVP scope and risks."
    assert parsed.markdown == parsed.plain_text
    assert parsed.parser.name == "utf8_text"
    assert parsed.quality.block_count == 2
    assert [block.type for block in parsed.blocks] == ["heading", "paragraph"]
    assert parsed.blocks[0].text_span == {"start": 0, "end": 8}


def test_parse_pdf_prefers_docling_when_available(tmp_path, monkeypatch):
    path = tmp_path / "defense.pdf"
    path.write_bytes(b"fake pdf bytes")
    _, _, blocks = build_blocks_from_output([{"type": "paragraph", "text": "Docling text", "markdown": "Docling text"}])

    def fake_parse_docling_document(path: Path) -> ParsedDocument:
        return ParsedDocument(
            plain_text="Docling text",
            markdown="Docling text",
            blocks=blocks,
            parser=ParserInfo(name="docling", version="test"),
            quality=ParseQuality(char_count=12, block_count=1),
        )

    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: True)
    monkeypatch.setattr(docling_parser, "parse_docling_document", fake_parse_docling_document)

    parsed = parse_file_to_document(path)

    assert parsed.parser.name == "docling"
    assert parsed.plain_text == "Docling text"


def test_parse_docx_extracts_paragraphs_and_tables(tmp_path, monkeypatch):
    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: False)
    path = tmp_path / "defense.docx"
    document = DocxDocument()
    document.add_paragraph("Gate 2 investment defense")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Traction"
    table.cell(1, 1).text = "Growing"
    document.save(path)

    parsed = parse_file(path)

    assert "Gate 2 investment defense" in parsed
    assert "| Metric | Value |" in parsed
    assert "| Traction | Growing |" in parsed


def test_parse_docx_returns_markdown_tables_and_quality(tmp_path, monkeypatch):
    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: False)
    path = tmp_path / "defense.docx"
    document = DocxDocument()
    document.add_heading("Gate 2 investment defense", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Traction"
    table.cell(1, 1).text = "Growing"
    document.save(path)

    parsed = parse_file_to_document(path)

    assert parsed.parser.name == "ooxml-docx"
    assert parsed.quality.table_count == 1
    assert [block.type for block in parsed.blocks] == ["heading", "table"]
    assert "| Metric | Value |" in parsed.markdown
    assert "| Traction | Growing |" in parsed.markdown


def test_parse_docx_preserves_structured_markdown_from_word_xml(tmp_path, monkeypatch):
    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: False)
    path = tmp_path / "structured.docx"
    _write_structured_docx(path)

    parsed = parse_file_to_document(path)

    expected_markdown = [
        "# Proposal",
        "1. First metric",
        "2. Second metric<br>with continuation",
        "| Metric | Value |",
        "| --- | --- |",
        "| Baseline | 10% \\| 20% |",
        "| Evidence | [external source](https://example.com/report) |",
        "| Spanning note |  |",
    ]
    for expected in expected_markdown:
        assert expected in parsed.plain_text
        assert expected in parsed.markdown
    assert parsed.plain_text == parsed.markdown
    assert parsed.quality.table_count == 1


def test_parse_docx_ignores_word_revision_and_field_nodes(tmp_path, monkeypatch):
    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: False)
    path = tmp_path / "structured.docx"
    _write_structured_docx(path)

    parsed = parse_file(path)

    assert "Visible paragraph" in parsed
    assert "deleted text" not in parsed
    assert "moved text" not in parsed
    assert "PAGE" not in parsed


def test_parse_docx_uses_builtin_parser_when_docling_available(tmp_path, monkeypatch):
    path = tmp_path / "structured.docx"
    _write_structured_docx(path)

    def fail_if_docling_is_used(path: Path) -> ParsedDocument:
        raise AssertionError("DOCX should use the deterministic OOXML parser before Docling")

    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: True)
    monkeypatch.setattr(docling_parser, "parse_docling_document", fail_if_docling_is_used)

    parsed = parse_file_to_document(path)

    assert "# Proposal" in parsed.plain_text
    assert parsed.parser.options["source"] == "zipfile.ElementTree"


def test_parse_dotx_uses_docx_parser(tmp_path):
    path = tmp_path / "travel.dotx"
    document = DocxDocument()
    document.add_paragraph("Gate 2 travel benchmark original")
    document.save(path)
    _rewrite_main_content_type(path, from_value=DOCX_MAIN_CONTENT_TYPE, to_value=DOTX_MAIN_CONTENT_TYPE)

    parsed = parse_file(path)

    assert "Gate 2 travel benchmark original" in parsed


def test_parse_pdf_extracts_text_with_page_markers(tmp_path, monkeypatch):
    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: False)
    path = tmp_path / "defense.pdf"
    path.write_bytes(_pdf_with_text_pages(["Gate 2 MVP traction", "Second page risks"]))

    parsed = parse_file(path)

    assert "[Page 1]" in parsed
    assert "Gate 2 MVP traction" in parsed
    assert "[Page 2]" in parsed
    assert "Second page risks" in parsed


def test_parse_pdf_returns_page_blocks_and_quality(tmp_path, monkeypatch):
    monkeypatch.setattr(docling_parser, "is_docling_available", lambda: False)
    path = tmp_path / "defense.pdf"
    path.write_bytes(_pdf_with_text_pages(["Gate 2 MVP traction", "Second page risks"]))

    parsed = parse_file_to_document(path)

    assert parsed.parser.name == "pypdf"
    assert parsed.quality.page_count == 2
    assert parsed.quality.empty_pages == []
    assert [block.page for block in parsed.blocks] == [1, 2]
    assert parsed.blocks[0].type == "page"
    assert "<!-- page 1 -->" in parsed.markdown


def _pdf_with_text_pages(page_texts: list[str]) -> bytes:
    objects: list[tuple[int, bytes]] = []
    kids: list[str] = []
    font_object_number = 3 + (len(page_texts) * 2)

    for index, text in enumerate(page_texts):
        page_object_number = 3 + (index * 2)
        content_object_number = 4 + (index * 2)
        kids.append(f"{page_object_number} 0 R")
        safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET".encode("latin-1")
        objects.append(
            (
                content_object_number,
                b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
            )
        )
        objects.append(
            (
                page_object_number,
                (
                    f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                    f"/Resources << /Font << /F1 {font_object_number} 0 R >> >> "
                    f"/Contents {content_object_number} 0 R >>"
                ).encode("ascii"),
            )
        )

    objects.extend(
        [
            (1, b"<< /Type /Catalog /Pages 2 0 R >>"),
            (2, f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {len(kids)} >>".encode("ascii")),
            (font_object_number, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"),
        ]
    )
    objects.sort(key=lambda item: item[0])

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = {0: 0}
    for object_number, body in objects:
        offsets[object_number] = len(pdf)
        pdf.extend(f"{object_number} 0 obj\n".encode("ascii"))
        pdf.extend(body)
        pdf.extend(b"\nendobj\n")

    xref_position = len(pdf)
    max_object_number = max(offsets)
    pdf.extend(f"xref\n0 {max_object_number + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for object_number in range(1, max_object_number + 1):
        pdf.extend(f"{offsets.get(object_number, 0):010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer\n<< /Size {max_object_number + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_position}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(pdf)


def _write_structured_docx(path: Path) -> None:
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
</Types>
"""
    root_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    <w:p>
      <w:pPr><w:pStyle w:val="Heading1"/></w:pPr>
      <w:r><w:t>Proposal</w:t></w:r>
    </w:p>
    <w:p>
      <w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="7"/></w:numPr></w:pPr>
      <w:r><w:t>First metric</w:t></w:r>
    </w:p>
    <w:p>
      <w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="7"/></w:numPr></w:pPr>
      <w:r><w:t>Second metric</w:t><w:br/><w:t>with continuation</w:t></w:r>
    </w:p>
    <w:p>
      <w:r><w:t>Visible paragraph</w:t></w:r>
      <w:del><w:r><w:t>deleted text</w:t></w:r></w:del>
      <w:moveFrom><w:r><w:t>moved text</w:t></w:r></w:moveFrom>
      <w:r><w:instrText>PAGE</w:instrText></w:r>
    </w:p>
    <w:tbl>
      <w:tr>
        <w:tc><w:p><w:r><w:t>Metric</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>Value</w:t></w:r></w:p></w:tc>
      </w:tr>
      <w:tr>
        <w:tc><w:p><w:r><w:t>Baseline</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>10% | 20%</w:t></w:r></w:p></w:tc>
      </w:tr>
      <w:tr>
        <w:tc><w:p><w:r><w:t>Evidence</w:t></w:r></w:p></w:tc>
        <w:tc>
          <w:p>
            <w:hyperlink r:id="rId5"><w:r><w:t>external source</w:t></w:r></w:hyperlink>
          </w:p>
        </w:tc>
      </w:tr>
      <w:tr>
        <w:tc>
          <w:tcPr><w:gridSpan w:val="2"/></w:tcPr>
          <w:p><w:r><w:t>Spanning note</w:t></w:r></w:p>
        </w:tc>
      </w:tr>
    </w:tbl>
  </w:body>
</w:document>
"""
    styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/></w:style>
</w:styles>
"""
    numbering_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="3">
    <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/></w:lvl>
  </w:abstractNum>
  <w:num w:numId="7"><w:abstractNumId w:val="3"/></w:num>
</w:numbering>
"""
    relationships_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://example.com/report" TargetMode="External"/>
</Relationships>
"""
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", root_rels_xml)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", styles_xml)
        archive.writestr("word/numbering.xml", numbering_xml)
        archive.writestr("word/_rels/document.xml.rels", relationships_xml)


def _rewrite_main_content_type(path: Path, *, from_value: str, to_value: str) -> None:
    temp_path = path.with_suffix(".tmp")
    with ZipFile(path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(from_value.encode("utf-8"), to_value.encode("utf-8"))
            target.writestr(item, data)
    temp_path.replace(path)
